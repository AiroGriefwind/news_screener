import os
import json
import time
import re
from pathlib import Path
from typing import Dict, List, Optional
from openai import OpenAI
from docx import Document
import argparse
from datetime import datetime

class NewsClassifier:
    def __init__(self, api_key: str, rate_limit_delay: float = 2.0):
        """Initialize the News Classifier for Kimi AI"""
        self.client = OpenAI(
            api_key=api_key,
            base_url="https://api.moonshot.cn/v1"
        )
        self.rate_limit_delay = rate_limit_delay
        print("Kimi AI client initialized successfully")

    def is_new_metadata_format(self, text: str) -> bool:
        """Check if the text is a metadata line (has at least two '|' separators)"""
        if not text:
            return False
        return text.count('|') >= 2

    def extract_articles_from_docx(self, file_path: str) -> List[Dict[str, str]]:
        """
        Extract articles from DOCX using metadata-based detection.
        Preserves original paragraph structure.
        """
        doc = Document(file_path)
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        articles = []
        current_article = None
        i = 0
        n = len(paragraphs)
        
        while i < n:
            line = paragraphs[i]
            # Found a metadata line
            if self.is_new_metadata_format(line):
                # The line before metadata is the title
                title_line = paragraphs[i-1] if i >= 1 else ""
                # Save previous article if exists
                if current_article:
                    articles.append(current_article)
                # Start new article
                current_article = {
                    'title': title_line,
                    'metadata': line,
                    'content_paragraphs': [],  # Store as list of paragraphs
                    'section': self._extract_section_from_metadata(line)
                }
                i += 1  # Move past metadata line
                # Collect content paragraphs until next article title
                while i < n:
                    # Check if next line is metadata (meaning current line is title of next article)
                    if i + 1 < n and self.is_new_metadata_format(paragraphs[i + 1]):
                        break
                    current_article['content_paragraphs'].append(paragraphs[i])
                    i += 1
                # Also store as single string for analysis
                current_article['content'] = ' '.join(current_article['content_paragraphs'])
            else:
                i += 1
        
        # Don't forget the last article
        if current_article:
            articles.append(current_article)
        return articles

    def _extract_section_from_metadata(self, metadata: str) -> str:
        """Extract section info from metadata line"""
        parts = metadata.split('|')[0].strip().split()
        if len(parts) >= 3:
            return parts[-1]
        return "Unknown"

    def normalize_location(self, location: str) -> str:
        """Normalize location names to country names"""
        location_map = {
            # Americas
            '美国': 'United States', '美國': 'United States', 'America': 'United States',
            'Alaska': 'United States', '阿拉斯加': 'United States',
            '华盛顿': 'United States', '華盛頓': 'United States', 'Washington': 'United States',
            
            # Russia
            '俄罗斯': 'Russia', '俄羅斯': 'Russia', '莫斯科': 'Russia', 'Moscow': 'Russia',
            
            # Europe (including Ukraine-related terms)
            '欧洲': 'Europe', '歐洲': 'Europe', '欧盟': 'Europe', '歐盟': 'Europe',
            '乌克兰': 'Europe', '烏克蘭': 'Europe', 'Ukraine': 'Europe',
            '基辅': 'Europe', '基輔': 'Europe', 'Kiev': 'Europe', 'Kyiv': 'Europe',
            '德国': 'Europe', '德國': 'Europe', 'Germany': 'Europe',
            '法国': 'Europe', '法國': 'Europe', 'France': 'Europe',
            '英国': 'Europe', '英國': 'Europe', 'UK': 'Europe', 'Britain': 'Europe',
            '意大利': 'Europe', 'Italy': 'Europe',
            '西班牙': 'Europe', 'Spain': 'Europe',
            '葡萄牙': 'Europe', 'Portugal': 'Europe',
            '希腊': 'Europe', '希臘': 'Europe', 'Greece': 'Europe',
            '土耳其': 'Europe', 'Turkey': 'Europe',
            '塞尔维亚': 'Europe', '塞爾維亞': 'Europe', 'Serbia': 'Europe',
            '布鲁塞尔': 'Europe', '布魯塞爾': 'Europe', 'Brussels': 'Europe',
            
            # Middle East
            '以色列': 'Middle East', 'Israel': 'Middle East',
            '巴勒斯坦': 'Middle East', 'Palestine': 'Middle East',
            '加沙': 'Middle East', 'Gaza': 'Middle East',
            '埃及': 'Middle East', 'Egypt': 'Middle East',
            
            # Southeast Asia
            '新加坡': 'Southeast Asia', 'Singapore': 'Southeast Asia',
            '菲律宾': 'Southeast Asia', '菲律賓': 'Southeast Asia', 'Philippines': 'Southeast Asia',
            '马尼拉': 'Southeast Asia', '馬尼拉': 'Southeast Asia', 'Manila': 'Southeast Asia',
            
            # Japan
            '日本': 'Japan', '大阪': 'Japan', 'Osaka': 'Japan', '道頓堀': 'Japan',
            '东京': 'Japan', '東京': 'Japan', 'Tokyo': 'Japan',
            
            # Korea
            '韩国': 'Korea', '韓國': 'Korea', 'South Korea': 'Korea',
            '朝鲜': 'Korea', '朝鮮': 'Korea', 'North Korea': 'Korea',
            
            # China
            '中国': 'China', '中國': 'China', 'China': 'China',
        }
        
        # Check for exact matches first
        for key, value in location_map.items():
            if key.lower() == location.lower():
                return value
        
        # Then check for partial matches
        for key, value in location_map.items():
            if key.lower() in location.lower() or location.lower() in key.lower():
                return value
        
        return location  # Return original if no mapping found




    def analyze_article_with_kimi(self, article: Dict[str, str]) -> Dict:
        """Send article to Kimi AI for 5W1H analysis with improved topic key generation"""
        prompt = f"""
请分析这篇新闻文章的"硬新闻"特征，基于5W1H标准（谁、什么、何时、何地、为什么、如何）。

文章标题：{article['title']}
文章内容：{article['content'][:3000]}

请评估：
1. 谁(WHO)：是否清楚识别了关键人物/实体？(评分1-5)
2. 什么(WHAT)：是否清楚描述了具体事件/行动？(评分1-5)
3. 何时(WHEN)：是否提供了时间/日期信息？(评分1-5)
4. 何地(WHERE)：是否明确指定了地点？(评分1-5)
5. 为什么(WHY)：是否解释了原因/动机？(评分1-5)
6. 如何(HOW)：是否描述了方法/过程？(评分1-5)

额外要求：
- 生成一个1-2句的简要主题总结（topic_summary）。
- 提取新闻的主要地点（main_location），优先使用国家名而非城市名。
- 判断是否与科技公司相关（is_tech_news）：如果提到"Meta"、"Space X"、"Elon Musk"、"Apple"、"Google"、"Microsoft"、"Tesla"、"Amazon"、"OpenAI"、"ChatGPT"等，则为true。
- 生成一个简短但具体的主题键（topic_key），用于识别相同事件。例如：
  * 大阪火灾 -> "osaka_fire"
  * 乌克兰战争谈判 -> "ukraine_peace_talks" 
  * 以色列加沙军事行动 -> "israel_gaza_military"
  * 日本游客菲律宾遇害 -> "japanese_tourists_philippines_murder"
  确保相同事件的文章使用相同的topic_key。

请以以下JSON格式返回分析结果：
{{
  "is_hard_news": true/false,
  "overall_score": 0-30,
  "analysis": {{
    "who_score": 1-5,
    "what_score": 1-5,
    "when_score": 1-5,
    "where_score": 1-5,
    "why_score": 1-5,
    "how_score": 1-5
  }},
  "missing_elements": ["缺失的5W1H元素列表"],
  "recommendation": "分类说明的简要解释",
  "first_three_paragraphs_analysis": "分析前三段是否包含5W1H信息",
  "topic_summary": "1-2句主题总结",
  "main_location": "主要地点字符串",
  "is_tech_news": true/false,
  "topic_key": "简短具体主题键"
}}

判断"硬新闻"的标准：
- 总分 ≥ 20（满分30）
- 至少4个元素得分 ≥ 3
- 前三段包含关键信息（谁、什么、何时、何地）
"""
        
        try:
            messages = [
                {
                    "role": "system",
                    "content": "你是Kimi，由Moonshot AI提供的人工智能助手。你擅长中文和英文对话，专业分析新闻内容的5W1H结构。特别注意生成准确的topic_key来识别相同的新闻事件。"
                },
                {
                    "role": "user",
                    "content": prompt
                }
            ]
            completion = self.client.chat.completions.create(
                model="moonshot-v1-32k",
                messages=messages,
                temperature=0.3
            )
            response_text = completion.choices[0].message.content
            
            # Try to extract JSON from the response
            json_match = re.search(r'\{.*\}', response_text, re.DOTALL)
            if json_match:
                json_str = json_match.group()
                analysis = json.loads(json_str)
                # Normalize the location
                if 'main_location' in analysis:
                    analysis['main_location'] = self.normalize_location(analysis['main_location'])
            else:
                # Fallback if JSON extraction fails
                analysis = {
                    "is_hard_news": False,
                    "overall_score": 0,
                    "analysis": {
                        "who_score": 0, "what_score": 0, "when_score": 0,
                        "where_score": 0, "why_score": 0, "how_score": 0
                    },
                    "missing_elements": ["all"],
                    "recommendation": "API response parsing failed",
                    "first_three_paragraphs_analysis": "Unable to analyze",
                    "topic_summary": "Unable to summarize",
                    "main_location": "Others",
                    "is_tech_news": False,
                    "topic_key": "unknown"
                }
            return analysis
            
        except Exception as e:
            print(f"Error analyzing article '{article['title'][:30]}...': {str(e)}")
            return {
                "is_hard_news": False,
                "overall_score": 0,
                "analysis": {
                    "who_score": 0, "what_score": 0, "when_score": 0,
                    "where_score": 0, "why_score": 0, "how_score": 0
                },
                "missing_elements": ["all"],
                "recommendation": f"Error: {str(e)}",
                "first_three_paragraphs_analysis": "Error occurred",
                "topic_summary": "Error occurred",
                "main_location": "Others",
                "is_tech_news": False,
                "topic_key": "error"
            }

    def process_document(self, input_file_path: str) -> None:
        """Process entire document and save results"""
        print(f"Processing document: {input_file_path}")
        
        # Extract articles using improved metadata-based method
        print("Extracting articles using metadata detection...")
        articles = self.extract_articles_from_docx(input_file_path)
        print(f"Found {len(articles)} articles")
        
        # Display first few articles for verification
        print("\n=== EXTRACTION VERIFICATION ===")
        for i, article in enumerate(articles[:3], 1):
            print(f"\nArticle {i}:")
            print(f"Title: {article['title']}")
            print(f"Metadata: {article['metadata']}")
            print(f"Content preview: {article['content'][:100]}...")
            print(f"Section: {article['section']}")
        
        # Process each article
        results = []
        for i, article in enumerate(articles, 1):
            print(f"Analyzing article {i}/{len(articles)}: {article['title'][:50]}...")
            # Analyze with Kimi AI
            analysis = self.analyze_article_with_kimi(article)
            # Combine article data with analysis
            result = {
                "article_info": {
                    "title": article["title"],
                    "metadata": article["metadata"],
                    "section": article["section"],
                    "content_length": len(article["content"]),
                    "content_preview": article["content"][:200] + "..." if len(article["content"]) > 200 else article["content"],
                    "content_paragraphs": article["content_paragraphs"],  # Preserve original paragraphs
                    "full_content": article["content"]
                },
                "analysis": analysis,
                "timestamp": datetime.now().isoformat(),
                "article_index": i
            }
            results.append(result)
            # Rate limiting
            time.sleep(self.rate_limit_delay)
        
        # Save all results to JSON
        self.save_results(input_file_path, results)
        
        # Deduplicate and select highest scoring per topic_key
        selected_results = self.deduplicate_results(results)
        
        # Create output DOCX
        self.create_output_docx(input_file_path, results, selected_results)
        
        # Print summary for all and selected
        self.print_summary(results, selected_results)

    def deduplicate_results(self, results: List[Dict]) -> List[Dict]:
        """Deduplicate by topic_key, keeping the one with highest overall_score"""
        from collections import defaultdict
        
        groups = defaultdict(list)
        for res in results:
            topic_key = res["analysis"].get("topic_key", f"unknown_{res['article_index']}")
            groups[topic_key].append(res)
        
        selected = []
        for key, group in groups.items():
            if group:
                # Sort by overall_score descending, then by article_index ascending
                group.sort(key=lambda x: (-x["analysis"]["overall_score"], x["article_index"]))
                best = group[0]
                
                # Log deduplication if multiple articles in group
                if len(group) > 1:
                    print(f"\nDuplicates found for topic '{key}':")
                    for item in group:
                        print(f"  - {item['article_info']['title'][:50]}... (Score: {item['analysis']['overall_score']})")
                    print(f"  Selected: {best['article_info']['title'][:50]}...")
                
                selected.append(best)
        
        return selected

    def create_output_docx(self, input_file_path: str, all_results: List[Dict], selected_results: List[Dict]) -> None:
        """Create a DOCX with full list summary and selected articles in specified order"""
        input_path = Path(input_file_path)
        output_file = input_path.parent / f"{input_path.stem}_selected_articles.docx"
        doc = Document()

        # Define location order with proper mapping
        location_order = {
            'United States': 0,
            'Russia': 1,
            'Europe': 2,
            'Middle East': 3,
            'Southeast Asia': 4,
            'Japan': 5,
            'Korea': 6,
            'China': 7,
            'Others': 8,
            'Tech': 9  # Tech news goes last
        }

        def get_sort_key(result):
            analysis = result["analysis"]
            if analysis.get("is_tech_news", False):
                return (9, result["article_index"])  # Tech news last
            location = analysis.get("main_location", "Others")
            # Apply normalization again to ensure consistency
            location = self.normalize_location(location)
            order = location_order.get(location, 8)  # Default to Others
            return (order, result["article_index"])

        # Apply normalization to all selected results for consistency
        for result in selected_results:
            if 'main_location' in result["analysis"]:
                result["analysis"]["main_location"] = self.normalize_location(result["analysis"]["main_location"])

        # Sort selected results by location order
        sorted_selected = sorted(selected_results, key=get_sort_key)

        # Debug: Print the sorting order
        print("\n=== SORTING DEBUG ===")
        for res in sorted_selected:
            analysis = res["analysis"]
            location = "Tech News" if analysis.get("is_tech_news", False) else analysis.get("main_location", "Others")
            title = res["article_info"]["title"][:50]
            print(f"{location}: {title}...")

        # Add full list summary
        doc.add_heading("Full Hard News Summary", level=1)
        hard_news = [r for r in all_results if r["analysis"]["is_hard_news"]]
        doc.add_paragraph(f"Total Hard News: {len(hard_news)}")
        doc.add_paragraph(f"Selected Articles (after deduplication): {len(selected_results)}")
        
        for res in hard_news:
            title = res["article_info"]["title"]
            score = res["analysis"]["overall_score"]
            doc.add_paragraph(f"- {title} (Score: {score}/30)")

        # Add selected articles with preserved formatting
        doc.add_heading("Selected Articles (Deduplicated and Ordered by Location)", level=1)
        
        # Group articles by location first
        from collections import defaultdict
        location_groups = defaultdict(list)
        
        for res in sorted_selected:
            analysis = res["analysis"]
            display_location = "Tech News" if analysis.get("is_tech_news", False) else analysis.get("main_location", "Others")
            location_groups[display_location].append(res)
        
        # Process each location group in order
        location_display_order = ['United States', 'Russia', 'Europe', 'Middle East', 'Southeast Asia', 'Japan', 'Korea', 'China', 'Others', 'Tech News']
        
        for location in location_display_order:
            if location in location_groups:
                # Add location header
                doc.add_heading(f"=== {location} ===", level=1)
                
                # Add all articles for this location
                for res in location_groups[location]:
                    info = res["article_info"]
                    analysis = res["analysis"]
                    
                    # Add article
                    doc.add_heading(info["title"], level=2)
                    doc.add_paragraph(f"{info['metadata']}")
                    #doc.add_paragraph(f"Section: {info['section']}")
                    #doc.add_paragraph(f"Score: {analysis['overall_score']}/30")
                    #doc.add_paragraph(f"Topic Summary: {analysis.get('topic_summary', 'N/A')}")
                    #doc.add_paragraph(f"Topic Key: {analysis.get('topic_key', 'N/A')}")
                    
                    # Add content preserving original paragraph structure
                    #doc.add_paragraph("Content:")
                    for paragraph in info.get("content_paragraphs", [info["full_content"]]):
                        if paragraph.strip():
                            doc.add_paragraph(paragraph, style='Normal')

        doc.save(output_file)
        print(f"Output DOCX saved to: {output_file}")

    def save_results(self, input_file_path: str, results: List[Dict]) -> None:
        """Save analysis results to JSON file"""
        input_path = Path(input_file_path)
        output_file = input_path.parent / f"{input_path.stem}_kimi_analysis.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(results, f, ensure_ascii=False, indent=2)
        print(f"Results saved to: {output_file}")

    def print_summary(self, all_results: List[Dict], selected_results: List[Dict]) -> None:
        """Print analysis summary for all and selected"""
        total_articles = len(all_results)
        hard_news_count = sum(1 for r in all_results if r["analysis"]["is_hard_news"])
        selected_count = len(selected_results)
        
        print("\n" + "="*50)
        print("KIMI AI 新闻分析总结")
        print("="*50)
        print(f"总分析文章数: {total_articles}")
        print(f"硬新闻文章数: {hard_news_count}")
        print(f"软新闻文章数: {total_articles - hard_news_count}")
        print(f"硬新闻比例: {(hard_news_count/total_articles)*100:.1f}%")
        print(f"选定文章数 (去重后): {selected_count}")
        print(f"去重减少数量: {hard_news_count - selected_count}")
        
        print(f"\n选定文章列表 (按地点排序):")
        for result in selected_results:
            score = result["analysis"]["overall_score"]
            title = result["article_info"]["title"]
            location = 'Tech' if result["analysis"]["is_tech_news"] else result["analysis"]["main_location"]
            topic_key = result["analysis"].get("topic_key", "N/A")
            print(f" - {title[:50]}... (Score: {score}/30, Location: {location}, Topic: {topic_key})")

def main():
    parser = argparse.ArgumentParser(description="使用Kimi AI分类新闻文章 - 改进版")
    parser.add_argument("input_file", help="输入.docx文件路径")
    parser.add_argument("--api-key", required=True, help="Kimi AI API密钥")
    parser.add_argument("--delay", type=float, default=2.0, help="API调用间隔（秒）")
    
    args = parser.parse_args()
    
    # Validate input file
    if not os.path.exists(args.input_file):
        print(f"错误: 找不到输入文件 '{args.input_file}'")
        return
    if not args.input_file.endswith('.docx'):
        print("错误: 输入文件必须是.docx格式")
        return
    
    # Initialize classifier and process
    classifier = NewsClassifier(args.api_key, args.delay)
    classifier.process_document(args.input_file)

if __name__ == "__main__":
    main()
